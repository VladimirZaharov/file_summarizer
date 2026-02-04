"""
File parsers module for different document types.
Modular design allows easy addition of new file type parsers.
"""
from pathlib import Path
from typing import Protocol, Dict, Type
import os


class FileParser(Protocol):
    """Protocol for file parsers"""
    
    @staticmethod
    def can_parse(file_path: Path) -> bool:
        """Check if this parser can handle the file"""
        ...
    
    @staticmethod
    def parse(file_path: Path) -> str:
        """Parse file and return text content"""
        ...


class TextFileParser:
    """Parser for plain text files"""
    
    @staticmethod
    def can_parse(file_path: Path) -> bool:
        return file_path.suffix.lower() in {'.txt', '.md', '.csv', '.log'}
    
    @staticmethod
    def parse(file_path: Path) -> str:
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, 'rb') as f:
                return f.read().decode('utf-8', errors='ignore')
        except Exception as e:
            return f"Error reading text file: {str(e)}"


class PDFParser:
    """Parser for PDF files"""
    
    @staticmethod
    def can_parse(file_path: Path) -> bool:
        return file_path.suffix.lower() == '.pdf'
    
    @staticmethod
    def parse(file_path: Path) -> str:
        try:
            import PyPDF2
            
            text = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text.append(f"--- Page {page_num + 1} ---\n{page_text}")
            
            return '\n\n'.join(text) if text else "No text extracted from PDF"
        
        except ImportError:
            return "PyPDF2 not installed. Install with: pip install PyPDF2"
        except Exception as e:
            return f"Error reading PDF file: {str(e)}"


class DocxParser:
    """Parser for DOCX files"""
    
    @staticmethod
    def can_parse(file_path: Path) -> bool:
        return file_path.suffix.lower() in {'.docx', '.doc'}
    
    @staticmethod
    def parse(file_path: Path) -> str:
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text.append(row_text)
            
            return '\n\n'.join(text) if text else "No text extracted from DOCX"
        
        except ImportError:
            return "python-docx not installed. Install with: pip install python-docx"
        except Exception as e:
            # Try alternative method for .doc files
            if file_path.suffix.lower() == '.doc':
                return "Legacy .doc files require additional libraries (e.g., antiword or LibreOffice)"
            return f"Error reading DOCX file: {str(e)}"


class ExcelParser:
    """Parser for Excel files"""
    
    @staticmethod
    def can_parse(file_path: Path) -> bool:
        return file_path.suffix.lower() in {'.xlsx', '.xls'}
    
    @staticmethod
    def parse(file_path: Path) -> str:
        try:
            import openpyxl
            
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text.append(f"=== Sheet: {sheet_name} ===")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                    if row_text.strip():
                        text.append(row_text)
                
                text.append("")  # Empty line between sheets
            
            return '\n'.join(text) if text else "No data extracted from Excel"
        
        except ImportError:
            return "openpyxl not installed. Install with: pip install openpyxl"
        except Exception as e:
            return f"Error reading Excel file: {str(e)}"


class HTMLParser:
    """Parser for HTML files"""
    
    @staticmethod
    def can_parse(file_path: Path) -> bool:
        return file_path.suffix.lower() in {'.html', '.htm'}
    
    @staticmethod
    def parse(file_path: Path) -> str:
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(['script', 'style']):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text if text else "No text extracted from HTML"
        
        except ImportError:
            return "beautifulsoup4 not installed. Install with: pip install beautifulsoup4"
        except Exception as e:
            return f"Error reading HTML file: {str(e)}"


class RTFParser:
    """Parser for RTF files"""
    
    @staticmethod
    def can_parse(file_path: Path) -> bool:
        return file_path.suffix.lower() == '.rtf'
    
    @staticmethod
    def parse(file_path: Path) -> str:
        try:
            from striprtf.striprtf import rtf_to_text
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            
            text = rtf_to_text(rtf_content)
            return text if text.strip() else "No text extracted from RTF"
        
        except ImportError:
            return "striprtf not installed. Install with: pip install striprtf"
        except Exception as e:
            return f"Error reading RTF file: {str(e)}"


class FileParserFactory:
    """Factory class to get appropriate parser for a file"""
    
    # Register all parsers here
    PARSERS: list[Type] = [
        TextFileParser,
        PDFParser,
        DocxParser,
        ExcelParser,
        HTMLParser,
        RTFParser,
    ]
    
    @classmethod
    def get_parser(cls, file_path: Path) -> Type:
        """Get appropriate parser for file"""
        for parser in cls.PARSERS:
            if parser.can_parse(file_path):
                return parser
        
        # Default to text parser for unknown types
        return TextFileParser
    
    @classmethod
    def parse_file(cls, file_path: Path) -> Dict[str, str]:
        """
        Parse a file and return its content.
        
        Returns:
            Dict with 'filename', 'type', and 'content'
        """
        parser = cls.get_parser(file_path)
        content = parser.parse(file_path)
        
        return {
            'filename': file_path.name,
            'type': file_path.suffix.lower(),
            'content': content,
            'size': file_path.stat().st_size
        }
    
    @classmethod
    def register_parser(cls, parser_class: Type):
        """Register a new parser dynamically"""
        if parser_class not in cls.PARSERS:
            cls.PARSERS.insert(0, parser_class)  # Insert at beginning for priority
